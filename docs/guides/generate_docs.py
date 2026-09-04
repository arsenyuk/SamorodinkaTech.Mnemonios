#!/usr/bin/env python3
"""Генератор документации ЕДИН в формате .docx."""

import platform
import sys

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ─────────────────────────── Общие утилиты ───────────────────────────


def setup_page(doc: Document, size: str = "A4") -> None:
    section = doc.sections[0]
    if size == "A4":
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
        section.top_margin = section.bottom_margin = Cm(2.54)
        section.left_margin = section.right_margin = Cm(3.18)
    else:
        from docx.shared import Inches
        section.page_width, section.height = Inches(8.5), Inches(11.0)
        section.top_margin = section.bottom_margin = Inches(1.0)
        section.left_margin = section.right_margin = Inches(1.25)
    section.orientation = WD_ORIENTATION.PORTRAIT


def tune_styles(doc: Document) -> None:
    body = doc.styles["Normal"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_after = Pt(6)

    title = doc.styles["Title"]
    title.font.name = "Calibri Light"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri Light"
    subtitle.font.size = Pt(16)
    subtitle.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    for n, size in [(1, 18), (2, 14), (3, 12)]:
        s = doc.styles[f"Heading {n}"]
        s.font.name = "Calibri Light"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        s.paragraph_format.space_before = Pt(14 - 2 * n)
        s.paragraph_format.space_after = Pt(4)


def patch_theme(doc: Document) -> None:
    cjk = {"Windows": "Microsoft YaHei", "Darwin": "PingFang SC"}.get(
        platform.system(), "Noto Sans CJK SC"
    )
    theme_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"
    try:
        theme_part = doc.part.part_related_by(theme_rel)
    except KeyError:
        return
    from lxml import etree
    theme_xml = etree.fromstring(theme_part.blob)
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    for latin in theme_xml.xpath(
        "//a:majorFont/a:latin | //a:minorFont/a:latin", namespaces=ns
    ):
        latin.set("typeface", "Calibri")
    for font in theme_xml.xpath(
        "//a:majorFont/a:font | //a:minorFont/a:font", namespaces=ns
    ):
        if font.get("script", "") in ("Hans", "Hant", "Jpan", "Hang"):
            font.set("typeface", cjk)
    theme_part._blob = etree.tostring(
        theme_xml, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    for style in doc.styles:
        if hasattr(style, "font") and style.font.name == "Courier":
            style.font.name = "Courier New"


def add_cover(doc: Document, title: str, subtitle: str = "", date: str = "") -> None:
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(subtitle, style="Subtitle")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(10):
        doc.add_paragraph()
    if date:
        p = doc.add_paragraph(date)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_toc(doc: Document) -> None:
    # Включаем автобновление полей при открытии в Word
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_text = OxmlElement("w:t")
    fld_char_text.text = "Содержание будет обновлено при открытии документа."
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    for el in (fld_char_begin, instr_text, fld_char_separate, fld_char_text, fld_char_end):
        run._r.append(el)


def add_page_number(paragraph: OxmlElement) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_footer(doc: Document, text: str = "") -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        footer.add_run(text + "  |  ")
    add_page_number(footer)


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, name in enumerate(header):
        hdr[i].text = name
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            cells[c_idx].text = str(value)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_code_block(doc: Document, code: str) -> None:
    for line in code.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def add_callout(doc: Document, text: str, color: str = "FFF4CE") -> None:
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)


def new_doc() -> Document:
    doc = Document()
    setup_page(doc)
    patch_theme(doc)
    tune_styles(doc)
    return doc


# ─────────────────────── Руководство администратора ───────────────────────


def create_admin_guide() -> Document:
    doc = new_doc()
    add_cover(doc, "ЕДИН", "Руководство администратора", "2026")
    doc.add_page_break()
    add_toc(doc)
    doc.add_page_break()
    add_footer(doc, "ЕДИН — Руководство администратора")

    # 1. Введение
    doc.add_heading("1. Введение", level=1)
    doc.add_paragraph(
        "Настоящее руководство предназначено для системных администраторов, "
        "ответственных за развёртывание, настройку и эксплуатацию сервиса "
        "ЕДИН (Master Person Index)."
    )
    doc.add_paragraph(
        "Сервис ЕДИН обеспечивает единую идентификацию физических лиц для "
        "корпоративных информационных систем. Система построена на .NET 10, "
        "использует PostgreSQL 16 и развёртывается через Docker Compose."
    )
    doc.add_paragraph(
        "Документ охватывает: системные требования, развёртывание, конфигурацию, "
        "управление сервисами, мониторинг, резервное копирование и безопасность."
    )

    # 2. Системные требования
    doc.add_heading("2. Системные требования", level=1)
    add_table(doc,
        ["Компонент", "Требование"],
        [
            ["ОС", "Linux (рекомендуется), macOS, Windows с Docker Desktop"],
            ["Docker", "20.10+ (с поддержкой Compose V2)"],
            ["Docker Compose", "2.0+"],
            ["RAM", "минимум 2 ГБ (рекомендуется 4 ГБ)"],
            ["Диск", "минимум 10 ГБ свободного места"],
            ["Сеть", "порты 5000 (API), 5003 (Steward), 5432 (PostgreSQL)"],
        ],
    )

    # 3. Развёртывание
    doc.add_heading("3. Развёртывание", level=1)

    doc.add_heading("3.1. Подготовка сервера", level=2)
    add_numbered(doc, "Установить Docker и Docker Compose.")
    add_numbered(doc, "Убедиться в доступности портов 5000, 5003, 5432.")
    add_numbered(doc, "Скопировать репозиторий на сервер.")

    doc.add_heading("3.2. Настройка переменных окружения", level=2)
    doc.add_paragraph(
        "Создайте файл .env в корне репозитория на основе .env.example:"
    )
    add_code_block(doc, "cp .env.example .env")
    doc.add_paragraph("Основные переменные:")
    add_table(doc,
        ["Переменная", "Значение по умолчанию", "Описание"],
        [
            ["POSTGRES_DB", "mnemonios", "Имя базы данных"],
            ["POSTGRES_USER", "mnemonios", "Пользователь PostgreSQL"],
            ["POSTGRES_PASSWORD", "mnemonios_dev", "Пароль PostgreSQL ( изменить в продакшене! )"],
            ["POSTGRES_PORT", "5432", "Внешний порт PostgreSQL"],
            ["API_PORT", "5000", "Внешний порт API"],
            ["STEWARD_PORT", "5003", "Внешний порт АРМ Стюарда"],
            ["HMAC_KEY", "(обязательно)", "Секретный ключ HMAC-SHA256 (минимум 32 символа)"],
            ["ASPNETCORE_ENVIRONMENT", "Production", "Среда выполнения"],
        ],
    )
    add_callout(doc, "Важно: HMAC_KEY — секретный ключ. Не храните его в git. Задайте уникальное значение минимум 32 символа.")

    doc.add_heading("3.3. Инициализация базы данных", level=2)
    doc.add_paragraph("Поднять PostgreSQL и применить каноническую схему:")
    add_code_block(doc, (
        "docker compose up -d postgres\n"
        "sleep 5\n"
        "docker exec -i mnemonios-postgres psql -U mnemonios -d mnemonios "
        "< tools/db/01_schema.sql"
    ))

    doc.add_heading("3.4. Запуск всех сервисов", level=2)
    add_code_block(doc, "docker compose up -d")
    doc.add_paragraph("Сервисы:")
    add_table(doc,
        ["Сервис", "Контейнер", "Назначение"],
        [
            ["postgres", "mnemonios-postgres", "База данных PostgreSQL 16"],
            ["api", "mnemonios-api", "REST API (эндпоинты идентификации)"],
            ["worker", "mnemonios-worker", "Фоновые задачи (реконсилизация)"],
            ["steward", "mnemonios-steward", "АРМ Стюарда (Razor Pages)"],
        ],
    )

    doc.add_heading("3.5. Проверка работоспособности", level=2)
    add_numbered(doc, "PostgreSQL: docker exec mnemonios-postgres pg_isready")
    add_numbered(doc, "API: http://<host>:5000/swagger")
    add_numbered(doc, "Steward: http://<host>:5003")
    add_numbered(doc, "Логи: docker logs mnemonios-api --tail 20")

    # 4. Конфигурация
    doc.add_heading("4. Конфигурация", level=1)

    doc.add_heading("4.1. AppSettings — Api", level=2)
    doc.add_paragraph("Основные секции в src/Applications/Api/appsettings.json:")
    add_table(doc,
        ["Секция", "Ключ", "Назначение"],
        [
            ["Serilog", "MinimumLevel.Default", "Уровень логирования (Information)"],
            ["Serilog", "WriteTo[].Args.path", "Путь к лог-файлам: logs/app-{yyyyMMddHH}.log"],
            ["ConnectionStrings", "DefaultConnection", "Строка подключения к PostgreSQL"],
            ["HmacSettings", "Key", "Секретный ключ HMAC ( из .env )"],
            ["HmacSettings", "DulSeparator", "Разделитель полей ДУЛ (по умолчанию | )"],
        ],
    )

    doc.add_heading("4.2. AppSettings — Worker", level=2)
    doc.add_paragraph("Конфигурация фоновых задач в src/Applications/Worker/appsettings.json:")
    add_code_block(doc, (
        '{\n'
        '  "Worker": {\n'
        '    "Tasks": [\n'
        '      {\n'
        '        "Id": "reconcile-cessations",\n'
        '        "CronExpression": "0 * * * *",\n'
        '        "Enabled": true,\n'
        '        "TimeoutSeconds": 300,\n'
        '        "RetryIntervalMinutes": 5\n'
        '      }\n'
        '    ]\n'
        '  }\n'
        '}'
    ))
    add_table(doc,
        ["Параметр", "Значение по умолчанию", "Описание"],
        [
            ["CronExpression", "0 * * * *", "Расписание: каждый час на 0-й минуте"],
            ["Enabled", "true", "Включена ли задача"],
            ["TimeoutSeconds", "300", "Таймаут выполнения (сек)"],
            ["RetryIntervalMinutes", "5", "Интервал повтора при ошибке (мин)"],
        ],
    )
    doc.add_paragraph("Примеры расписания:")
    add_bullet(doc, '"*/5 * * * *" — каждые 5 минут')
    add_bullet(doc, '"*/15 * * * *" — каждые 15 минут')
    add_bullet(doc, '"0 */2 * * *" — каждые 2 часа')

    doc.add_heading("4.3. Настройка часового пояса", level=2)
    doc.add_paragraph(
        "Контейнер Steward работает в часовом поясе Europe/Moscow (переменная TZ). "
        "Для изменения — обновите переменную TZ в docker-compose.yml."
    )

    # 5. Управление сервисами
    doc.add_heading("5. Управление сервисами", level=1)

    doc.add_heading("5.1. Остановка и запуск", level=2)
    add_code_block(doc, (
        "# Остановка всех сервисов\n"
        "docker compose down\n\n"
        "# Запуск всех сервисов\n"
        "docker compose up -d\n\n"
        "# Запуск только API\n"
        "docker compose up -d api\n\n"
        "# Запуск только Worker\n"
        "docker compose up -d worker\n\n"
        "# Запуск только Steward\n"
        "docker compose up -d steward"
    ))

    doc.add_heading("5.2. Пересборка при изменении кода", level=2)
    add_callout(doc, "Критически важно: при изменении исходного кода используйте docker compose build --no-cache, а не просто docker compose up -d. Без флага --no-cache Docker использует кэшированный слой COPY src/ src/ и контейнер запускается со старым кодом.")
    add_code_block(doc, (
        "# Пересборка и запуск API\n"
        "docker compose build --no-cache api\n"
        "docker compose up -d api\n\n"
        "# Пересборка и запуск Steward\n"
        "docker compose build --no-cache steward\n"
        "docker compose up -d steward\n\n"
        "# Пересборка и запуск Worker\n"
        "docker compose build --no-cache worker\n"
        "docker compose up -d worker"
    ))

    doc.add_heading("5.3. Просмотр логов", level=2)
    add_code_block(doc, (
        "# Логи API (последние 100 строк)\n"
        "docker logs mnemonios-api --tail 100\n\n"
        "# Логи Worker\n"
        "docker logs mnemonios-worker --tail 100\n\n"
        "# Логи Steward\n"
        "docker logs mnemonios-steward --tail 100\n\n"
        "# Непрерывный мониторинг\n"
        "docker logs -f mnemonios-api"
    ))

    doc.add_heading("5.4. Проверка обновления", level=2)
    doc.add_paragraph("После пересборки проверьте timestamp DLL в контейнере:")
    add_code_block(doc, "docker exec mnemonios-api ls -la /app/Api.dll")
    doc.add_paragraph("Timestamp должен соответствовать времени пересборки.")

    # 6. Мониторинг и диагностика
    doc.add_heading("6. Мониторинг и диагностика", level=1)

    doc.add_heading("6.1. Логи приложения", level=2)
    doc.add_paragraph(
        "Логи хранятся в каталоге logs/ внутри контейнера. Формат имени файла: "
        "app-{yyyyMMddHH}.log (покатный лог по часам)."
    )
    doc.add_paragraph("Для доступа к лог-файлам:")
    add_code_block(doc, "docker exec mnemonios-api ls -la /app/logs/")

    doc.add_heading("6.2. Мониторинг PostgreSQL", level=2)
    add_code_block(doc, (
        "# Проверка доступности\n"
        "docker exec mnemonios-postgres pg_isready -U mnemonios\n\n"
        "# Подключение к psql\n"
        "docker exec -it mnemonios-postgres psql -U mnemonios -d mnemonios\n\n"
        "# Количество записей в основных таблицах\n"
        "SELECT 'persons' as t, count(*) FROM persons\n"
        "UNION ALL SELECT 'keys', count(*) FROM person_identification_keys\n"
        "UNION ALL SELECT 'ext_persons', count(*) FROM ext_persons\n"
        "UNION ALL SELECT 'review_queue', count(*) FROM person_review_queue;"
    ))

    doc.add_heading("6.3. Очередь на ручную обработку", level=2)
    doc.add_paragraph(
        "Записи в person_review_queue возникают при статусе Ambiguous. "
        "Worker автоматически обрабатывает отложенные отзывы и реконсилизацию, "
        "но конфликты Ambiguous требуют ручного разбора через АРМ Стюарда."
    )

    # 7. Резервное копирование
    doc.add_heading("7. Резервное копирование и восстановление", level=1)

    doc.add_heading("7.1. Бэкап PostgreSQL", level=2)
    add_code_block(doc, (
        "# Бэкап базы данных\n"
        "docker exec mnemonios-postgres pg_dump -U mnemonios mnemonios > backup_$(date +%Y%m%d).sql\n\n"
        "# Бэкап с сжатием\n"
        "docker exec mnemonios-postgres pg_dump -U mnemonios mnemonios | gzip > backup_$(date +%Y%m%d).sql.gz"
    ))

    doc.add_heading("7.2. Восстановление из бэкапа", level=2)
    add_code_block(doc, (
        "# Восстановление\n"
        "docker exec -i mnemonios-postgres psql -U mnemonios -d mnemonios < backup_20250101.sql"
    ))

    doc.add_heading("7.3. Полный сброс и пересоздание", level=2)
    add_callout(doc, "Внимание: операция уничтожает все данные!", "FFECEC")
    add_code_block(doc, (
        "# Сброс схемы и данных\n"
        "docker exec -i mnemonios-postgres psql -U mnemonios -d mnemonios < tools/db/00_reset_schema.sql"
    ))

    # 8. Безопасность
    doc.add_heading("8. Безопасность", level=1)

    doc.add_heading("8.1. HMAC-ключ", level=2)
    doc.add_paragraph(
        "HMAC-ключ используется для вычисления идентификационных хешей. "
        "Хранится в переменной окружения HMAC_KEY (файл .env). "
        "Запрещено хранить его в appsettings.json, git или других отслеживаемых файлах."
    )

    doc.add_heading("8.2. Сетевая изоляция", level=2)
    doc.add_paragraph(
        "Все контейнеры работают в bridge-сети mnemonios. "
        "API и Steward доступны через маппинг портов. "
        "PostgreSQL доступен только из внутренней сети (порт 5432 маппится только для разработки)."
    )

    doc.add_heading("8.3. Проверка CVE", level=2)
    doc.add_paragraph(
        "Регулярно проверяйте компоненты на известные уязвимости (CVE) через "
        "БДУ ФСТЭК, GitHub Security Advisories, Snyk. "
        "Скрипт проверки: tools/check_fstec.py"
    )

    doc.add_heading("8.4. Ограничение доступа", level=2)
    doc.add_paragraph(
        "Сервис предназначен для работы в корпоративном контуре. "
        "API не реализует аутентификацию (TD-002 в tech_debt.md). "
        "Ограничьте доступ к портам 5000, 5003 на уровне сети (firewall, VPN)."
    )

    # 9. Обновление системы
    doc.add_heading("9. Обновление системы", level=1)
    add_numbered(doc, "Получить последние изменения: git pull origin develop")
    add_numbered(doc, "Проверить миграции БД (если изменилась 01_schema.sql — применить вручную)")
    add_numbered(doc, "Пересобрать контейнеры: docker compose build --no-cache api worker steward")
    add_numbered(doc, "Запустить: docker compose up -d")
    add_numbered(doc, "Проверить логи: docker logs -f mnemonios-api")

    doc.add_paragraph(
        "При несоответствии EF-модели и БД — приоритет у канонических SQL. "
        "Выравнивание через HasColumnName(...) в EF-конфигурациях."
    )

    # 10. Решение проблем
    doc.add_heading("10. Решение проблем", level=1)

    add_table(doc,
        ["Проблема", "Причина", "Решение"],
        [
            [
                "Контейнер запускается со старым кодом",
                "Docker использует кэш слоя COPY",
                "Использовать docker compose build --no-cache <service>",
            ],
            [
                "API не отвечает на /swagger",
                "Контейнер не запущен или упал",
                "docker logs mnemonios-api; docker compose up -d api",
            ],
            [
                "PostgreSQL не доступен",
                "Контейнер не прошёл health check",
                "docker compose restart postgres; проверить логи",
            ],
            [
                "Ошибка «HMAC key not configured»",
                "Не задана переменная HMAC_KEY",
                "Добавить HMAC_KEY в файл .env",
            ],
            [
                "Ошибки валидации ИНН/СНИЛС",
                "Неверная контрольная сумма",
                "Проверить корректность данных; дефект будет залогирован",
            ],
            [
                "Worker не обрабатывает задачи",
                "CronExpression неактивен или задача отключена",
                "Проверить Worker.Tasks[].Enabled и CronExpression",
            ],
        ],
    )

    doc.save("docs/guides/admin_guide.docx")
    print("  [OK] admin_guide.docx")


# ─────────────────────── Руководство пользователя ───────────────────────


def create_user_guide() -> Document:
    doc = new_doc()
    add_cover(doc, "ЕДИН", "Руководство пользователя", "2026")
    doc.add_page_break()
    add_toc(doc)
    doc.add_page_break()
    add_footer(doc, "ЕДИН — Руководство пользователя")

    # 1. Введение
    doc.add_heading("1. Введение", level=1)
    doc.add_paragraph(
        "Настоящее руководство предназначено для стюардов данных и операторов, "
        "работающих с АРМ Стюарда сервиса ЕДИН (Master Person Index)."
    )
    doc.add_paragraph(
        "АРМ Стюарда — веб-интерфейс для мониторинга и анализа данных MPI: "
        "очередь конфликтов, дефекты данных, история разборов, справочники."
    )

    # 2. Обзор АРМ Стюарда
    doc.add_heading("2. Обзор АРМ Стюарда", level=1)
    doc.add_paragraph("АРМ Стюарда доступен по адресу: http://<host>:5003")
    doc.add_paragraph("[вставить скриншот: общая навигация АРМ Стюарда]")
    doc.add_paragraph("Основные разделы навигации:")
    add_table(doc,
        ["Раздел", "Назначение"],
        [
            ["Очередь", "Конфликты (Ambiguous), требующие внимания стюарда"],
            ["Дефекты", "Мастер-записи с проблемами данных (ИНН, СНИЛС, ДУЛ)"],
            ["История", "Разрешённые конфликты (кто, когда, как)"],
            ["Классификатор ДУЛ", "Справочник кодов видов документов"],
            ["Маски URL", "Настройка URL-шаблонов для внешних ссылок"],
        ],
    )

    # 3. Очередь конфликтов
    doc.add_heading("3. Очередь конфликтов", level=1)

    doc.add_heading("3.1. Что такое конфликт", level=2)
    doc.add_paragraph(
        "Конфликт (статус Ambiguous) возникает, когда данные из внешней информационной "
        "системы совпадают с данными более чем одной записи в БД. "
        "Например, один и тот же паспорт зарегистрирован за двумя разными лицами — "
        "необходимо выяснить, одно это лицо или нет."
    )
    doc.add_paragraph("В очередь на обработку попадают только конфликты, не разрешённые автоматически.")

    doc.add_heading("3.2. Список ожидающих", level=2)
    doc.add_paragraph("Раздел «Очередь» отображает все записи с pending-статусом:")
    doc.add_paragraph("[вставить скриншот: список очереди конфликтов]")
    add_table(doc,
        ["Колонка", "Описание"],
        [
            ["Лицо A", "Существующая мастер-запись (PersonID)"],
            ["Лицо B", "Новая мастер-запись, присланная системой"],
            ["Ключ совпадения", "Тип ключа, по которому лица совпали (inn, dul, snils и т.д.)"],
            ["Ключ конфликта", "Тип ключа, по которому лица расходятся"],
            ["Дата", "Когда конфликт был создан"],
        ],
    )

    doc.add_heading("3.3. Детали конфликта", level=2)
    doc.add_paragraph("При переходе в детали конфликта отображается:")
    add_bullet(doc, "Таблица сравнения ключей (совпадают / конфликт / только у A / только у B)")
    add_bullet(doc, "Основные данные каждой персоны (PersonID, даты)")
    add_bullet(doc, "Связи с внешними системами")
    add_bullet(doc, "Идентификационные ключи")
    add_bullet(doc, "Дефекты данных")
    doc.add_paragraph("[вставить скриншот: детали конфликта — таблица сравнения ключей]")

    doc.add_heading("3.4. Действия стюарда", level=2)
    doc.add_paragraph("Два доступных действия:")
    add_bullet(doc, "Подтвердить слияние (B → A): все ключи, внешние ID и документы лица B переносятся в лицо A. Лицо B удаляется.")
    add_bullet(doc, "Отклонить: конфликт остаётся в очереди без изменений.")

    doc.add_heading("3.5. Автоматическое закрытие", level=2)
    doc.add_paragraph(
        "Если внешняя система сама исправляет данные (присылает новый resolve "
        "с корректными значениями), конфликт закрывается автоматически. "
        "Такие случаи отображаются в разделе «История» с типом resolution = auto_resolved."
    )

    # 4. Дефекты данных
    doc.add_heading("4. Дефекты данных", level=1)

    doc.add_paragraph(
        "Раздел «Дефекты» отображает все мастер-записи, имеющие проблемы с данными."
    )
    doc.add_paragraph("[вставить скриншот: список дефектов]")

    doc.add_heading("4.1. Типы дефектов", level=2)
    add_table(doc,
        ["Тип", "Описание", "Причина"],
        [
            ["invalid_inn", "Некорректный ИНН", "Неверная контрольная сумма ИНН"],
            ["invalid_snils", "Некорректный СНИЛС", "Неверная контрольная сумма СНИЛС"],
            ["dul_incomplete", "ДУЛ неполный", "Указана только серия или только номер документа"],
        ],
    )

    doc.add_heading("4.2. Детали дефекта", level=2)
    doc.add_paragraph("Для каждого дефекта отображаются:")
    add_bullet(doc, "Тип дефекта")
    add_bullet(doc, "Поле, вызвавшее дефект (inn, snils, dul_series, dul_number)")
    add_bullet(doc, "Сообщение об ошибке")
    doc.add_paragraph(
        "Дефекты логируются, но не блокируют создание записи. "
        "Это «толерантная обработка ошибок» — сервис фиксирует проблему, "
        "но не отказывается от добавления сведений о лице."
    )

    # 5. История конфликтов
    doc.add_heading("5. История конфликтов", level=1)
    doc.add_paragraph("Раздел «История» отображает все разрешённые конфликты.")
    doc.add_paragraph("[вставить скриншот: история конфликтов]")
    add_table(doc,
        ["Колонка", "Описание"],
        [
            ["Лицо A", "Идентификатор мастер-записи A"],
            ["Лицо B", "Идентификатор мастер-записи B"],
            ["Ключ совпадения", "Тип совпавшего ключа"],
            ["Ключ конфликта", "Тип конфликтующего ключа"],
            ["Результат", "Тип разрешения (auto_resolved)"],
            ["Разрешено кем", "source_system_id внешней системы"],
            ["Дата разрешения", "Дата и время разрешения"],
        ],
    )
    doc.add_paragraph(
        "При автоматическом разрешении (auto_resolved) в поле «Разрешено кем» "
        "указывается идентификатор внешней системы, которая исправила данные."
    )

    # 6. Карточка физического лица
    doc.add_heading("6. Карточка физического лица", level=1)
    doc.add_paragraph("Доступна по клику на PersonID из любого раздела.")
    doc.add_paragraph("[вставить скриншот: карточка физического лица]")

    doc.add_heading("6.1. Основные данные", level=2)
    add_bullet(doc, "PersonID (UUID)")
    add_bullet(doc, "Дата создания и последнего обновления")

    doc.add_heading("6.2. Связи с внешними системами", level=2)
    doc.add_paragraph(
        "Таблица всех external_id, привязанных к данному лицу. "
        "Если настроены Маски URL, отображаются кликабельные ссылки на внешние системы."
    )

    doc.add_heading("6.3. Идентификационные ключи", level=2)
    doc.add_paragraph(
        "Список HMAC-ключей (типы: inn, snils, dul, inn_fio, snils_fio, dul_fio). "
        "Ключи отображаются в усечённом виде (первые 16 символов). "
        "Полные ключи не отображаются в целях безопасности."
    )

    doc.add_heading("6.4. Документы (ДУЛ)", level=2)
    doc.add_paragraph(
        "Список документов, удостоверяющих личность, привязанных к лицу. "
        "Тип документа отображается по классификатору ДУЛ."
    )

    # 7. Staging-записи
    doc.add_heading("7. Staging-записи (ext_persons)", level=1)
    doc.add_paragraph(
        "Доступны из карточки физического лица. "
        "Отображают сырые данные всех запросов resolve, привязанных к данному лицу."
    )
    doc.add_paragraph("[вставить скриншот: staging-записи]")
    add_table(doc,
        ["Поле", "Описание"],
        [
            ["Source System", "Идентификатор внешней системы"],
            ["External Person ID", "Внешний идентификатор"],
            ["IP-адрес", "IP-адрес источника запроса"],
            ["Ключи HMAC", "HMAC-хеши всех вычисленных ключей"],
            ["Дефекты", "Дефекты, обнаруженные в данном запросе"],
        ],
    )

    # 8. Справочник ДУЛ
    doc.add_heading("8. Справочник ДУЛ", level=1)
    doc.add_paragraph(
        "Справочник кодов видов документов, удостоверяющих личность "
        "(Приказ ФНС России от 31.08.2020 № ЕД-7-14/617@)."
    )
    doc.add_paragraph("[вставить скриншот: справочник ДУЛ]")
    add_table(doc,
        ["Код", "Наименование документа"],
        [
            ["03", "Свидетельство о рождении"],
            ["07", "Военный билет"],
            ["08", "Временное удостоверение, выданное взамен военного билета"],
            ["10", "Паспорт иностранного гражданина"],
            ["11", "Свидетельство о рассмотрении ходатайства о признании беженцем"],
            ["12", "Вид на жительство в РФ"],
            ["13", "Удостоверение беженца"],
            ["15", "Разрешение на временное проживание в РФ"],
            ["18", "Свидетельство о предоставлении временного убежища"],
            ["21", "Паспорт гражданина РФ"],
            ["23", "Свидетельство о рождении (иностранное)"],
            ["24", "Удостоверение личности военнослужащего РФ"],
            ["91", "Иные документы"],
        ],
    )

    # 9. Маски URL
    doc.add_heading("9. Маски URL", level=1)
    doc.add_paragraph(
        "Маски URL позволяют настроить шаблоны ссылок на внешние информационные системы. "
        "При наличии маски ссылки отображаются в карточке физического лица."
    )
    doc.add_paragraph("[вставить скриншот: список масок URL]")

    doc.add_heading("9.1. Формат шаблона", level=2)
    doc.add_paragraph("Шаблон содержит плейсхолдеры:")
    add_bullet(doc, "{external_person_id} — внешний идентификатор лица")
    add_bullet(doc, "{source_system_id} — идентификатор системы")
    add_bullet(doc, "{organization_unit_key} — ключ организационной единицы")
    doc.add_paragraph("Пример: https://crm.example.com/persons/{external_person_id}")

    doc.add_heading("9.2. Создание и редактирование", level=2)
    doc.add_paragraph("[вставить скриншот: редактирование маски URL]")
    add_bullet(doc, "Укажите систему (sourceSystemId) и тип объекта (externalPersonType)")
    add_bullet(doc, "Введите URL-шаблон с плейсхолдерами")
    add_bullet(doc, "Сохраните")

    # 10. Рекомендации по работе
    doc.add_heading("10. Рекомендации по работе", level=1)

    doc.add_heading("10.1. Порядок обработки очереди", level=2)
    add_numbered(doc, "Откройте раздел «Очередь».")
    add_numbered(doc, "Изучите таблицу сравнения ключей.")
    add_numbered(doc, "Если ключи совпадают по ИНН или СНИЛС — скорее всего, одно лицо. Подтвердите слияние.")
    add_numbered(doc, "Если есть сомнения — отклоните и обратитесь к дополнительным источникам данных.")
    add_numbered(doc, "Проверьте историю разрешённых конфликтов для контекста.")

    doc.add_heading("10.2. Когда конфликт закрывается автоматически", level=2)
    doc.add_paragraph(
        "Конфликт закрывается автоматически, когда одна из внешних систем "
        "исправляет данные в последующем запросе resolve. "
        "Например: система CRM прислала ИНН с ошибкой → Ambiguous. "
        "Затем CRM исправила ИНН → конфликт закрыт, история сохранена."
    )
    doc.add_paragraph("Такие записи отображаются в «Истории» с resolution = auto_resolved.")

    doc.add_heading("10.3. Отличия auto-resolved от ручного разбора", level=2)
    add_table(doc,
        ["Параметр", "Auto-resolved", "Ручной разбор"],
        [
            ["Инициатор", "Внешняя система (исправление данных)", "Стюард"],
            ["Результат", "Конфликт закрыт автоматически", "Слияние или отклонение"],
            ["История", "resolved_by = source_system_id", "resolved_by = steward"],
        ],
    )

    doc.save("docs/guides/user_guide.docx")
    print("  [OK] user_guide.docx")


# ─────────────────────── Пояснительная записка ───────────────────────


def create_explanatory_note() -> Document:
    doc = new_doc()
    add_cover(doc, "ЕДИН", "Пояснительная записка", "2026")
    doc.add_page_break()
    add_toc(doc)
    doc.add_page_break()
    add_footer(doc, "ЕДИН — Пояснительная записка")

    # 1. Введение
    doc.add_heading("1. Введение", level=1)
    doc.add_paragraph(
        "Настоящая пояснительная записка описывает назначение, архитектуру, "
        "алгоритмы и нормативную базу сервиса ЕДИН (Master Person Index) — "
        "единого сервиса идентификации физических лиц для корпоративных "
        "информационных систем."
    )
    doc.add_paragraph(
        "Документ предназначен для технических специалистов, архитекторов, "
        "разработчиков и руководства проекта."
    )

    # 2. Назначение системы
    doc.add_heading("2. Назначение системы", level=1)

    doc.add_heading("2.1. Проблема", level=2)
    doc.add_paragraph(
        "Одна и та же персональная информация поступает из разных корпоративных "
        "систем (CRM, ERP, HR, бухгалтерия) в различных форматах, с ошибками "
        "и дублями. Без единого сервиса идентификации невозможно определить, "
        "относятся ли записи из разных систем к одному лицу."
    )
    doc.add_paragraph("Пример проблемы:")
    add_bullet(doc, "Система A: Иванов Иван Иванович, ИНН 7707083893")
    add_bullet(doc, "Система B: Иванов И.И., СНИЛС 12345678964")
    add_bullet(doc, "Система C: Иванов Иван, паспорт 4510 123456")
    doc.add_paragraph("Непонятно — одно это лицо или три разных.")

    doc.add_heading("2.2. Решение", level=2)
    doc.add_paragraph(
        "ЕДИН предоставляет единый PersonID (UUID), стабильный на протяжении "
        "всего жизненного цикла записи. Все записи из разных систем сопоставляются "
        "с единым мастер-записью."
    )

    doc.add_heading("2.3. Границы использования", level=2)
    add_table(doc,
        ["Категория", "Сервис предназначен для", "Сервис НЕ предназначен для"],
        [
            ["Идентификация", "Единое определение лица в корпоративном контуре", "Аутентификация или авторизация"],
            ["Данные", "HMAC-хеши ключей и связи с внешними системами", "Полный комплект ПДн"],
            ["Интеграция", "Стабильный PersonID для кросс-системной интеграции", "Замена карточек учёта"],
            ["Контроль", "Валидация ИНН, СНИЛС, ДУЛ", "Управление доступом"],
            ["Согласие", "Технический механизм удаления данных (cessation)", "Хранение согласий на обработку ПДн"],
        ],
    )

    # 3. Нормативная база
    doc.add_heading("3. Нормативная база", level=1)

    doc.add_heading("3.1. Федеральный закон № 152-ФЗ «О персональных данных»", level=2)
    add_table(doc,
        ["Статья", "Требование", "Реализация в ЕДИН"],
        [
            ["ст. 5 п. 1 ч. 1", "Обработка ПДн на основе согласия", "HMAC-хеши не являются ПДн. ФИО используются только в комбинированных ключах"],
            ["ст. 5 п. 5 ч. 5", "Уничтожение ПДн по требованию", "cessation + deferred cessation — транзакционное удаление"],
            ["ст. 14", "Информирование субъекта", "Аудит-лог: PersonId, HMAC-хеши, внешние ID (без ПДн)"],
            ["ст. 21", "Отзыв согласия", "POST /persons/cessation и POST /persons/cessation/deferred"],
        ],
    )

    doc.add_heading("3.2. Федеральный закон № 149-ФЗ «Об информации»", level=2)
    add_bullet(doc, "ст. 16 ч. 2: Защита информации — PostgreSQL 16, проверка CVE")
    add_bullet(doc, "ст. 16 ч. 3: Конфиденциальность — HMAC-SHA256, секретный ключ в .env")

    doc.add_heading("3.3. Постановление Правительства РФ № 1119", level=2)
    doc.add_paragraph(
        "Уровень защищённости: сервис обрабатывает ПДн в усечённом виде "
        "(пункт 11 Правил — данные не позволяют идентифицировать субъекта без внешних ключей)."
    )

    doc.add_heading("3.4. Приказ ФСТЭК России № 21", level=2)
    add_bullet(doc, "Аутентификация — вне scope (TD-002)")
    add_bullet(doc, "Регистрация событий безопасности — аудит-лог")
    add_bullet(doc, "Защита от вредоносного ПО — проверка CVE (tools/check_fstec.py)")

    doc.add_heading("3.5. ГОСТ Р 57580.1-2017", level=2)
    doc.add_paragraph(
        "Идентификация клиентов: ЕДИН обеспечивает единую идентификацию "
        "физических лиц для корпоративных систем. "
        "Управление рисками: валидация ИНН (10/12 цифр), СНИЛС (11 цифр), ДУЛ."
    )

    doc.add_heading("3.6. Постановление Правительства РФ № 1266", level=2)
    doc.add_paragraph(
        "Автоматизированная обработка ПДн: все операции через API с логированием. "
        "Ограничение обработки целями: только функция идентификации (MPI)."
    )

    # 4. Архитектура системы
    doc.add_heading("4. Архитектура системы", level=1)

    doc.add_heading("4.1. Архитектурные принципы", level=2)
    add_table(doc,
        ["Принцип", "Описание"],
        [
            ["Чистая архитектура", "Зависимости направлены внутрь: Api → Application → Domain ← Infrastructure"],
            ["Domain-Driven Design", "Моделирование через доменные сущности и ограниченные контексты"],
            ["SOLID", "Соблюдение всех пяти принципов"],
            ["Database-First", "Схема БД ведётся каноническим SQL-скриптом"],
        ],
    )

    doc.add_heading("4.2. Компоненты системы", level=2)
    add_table(doc,
        ["Компонент", "Технология", "Назначение"],
        [
            ["API", "ASP.NET Core Minimal API", "REST-эндпоинты идентификации"],
            ["Worker", ".NET Generic Host", "Фоновые задачи (реконсилизация)"],
            ["Steward", "ASP.NET Core Razor Pages", "АРМ Стюарда (мониторинг и разбор конфликтов)"],
            ["PostgreSQL", "PostgreSQL 16", "Хранение данных"],
            ["Docker", "Docker Compose", "Оркестрация контейнеров"],
        ],
    )

    doc.add_heading("4.3. Структура проекта", level=2)
    add_code_block(doc, (
        "SamorodinkaTech.Mnemonios/\n"
        "├── src/\n"
        "│   ├── Applications/\n"
        "│   │   ├── Api/           # REST API\n"
        "│   │   ├── Steward/      # АРМ Стюарда\n"
        "│   │   └── Worker/       # Фоновые задачи\n"
        "│   └── Core/\n"
        "│       ├── Domain/        # Сущности, интерфейсы, DTO\n"
        "│       ├── Infrastructure/# EF Core, сервисы, маппинг\n"
        "│       └── Common/        # Общие утилиты\n"
        "├── tests/\n"
        "├── docs/\n"
        "├── tools/db/              # SQL-скрипты схемы\n"
        "└── docker-compose.yml"
    ))

    doc.add_heading("4.4. Схема развёртывания", level=2)
    add_code_block(doc, (
        "┌─────────────────┐     ┌─────────────────┐\n"
        "│  API (:5000)    │     │  Steward (:5003)│\n"
        "└────────┬────────┘     └────────┬────────┘\n"
        "         │                       │\n"
        "         └───────────┬───────────┘\n"
        "                     │\n"
        "         ┌───────────▼───────────┐\n"
        "         │   PostgreSQL 16       │\n"
        "         │   (:5432)             │\n"
        "         └───────────────────────┘\n"
        "         ┌───────────────────────┐\n"
        "         │   Worker (cron)       │\n"
        "         └───────────────────────┘"
    ))

    # 5. Модель данных
    doc.add_heading("5. Модель данных", level=1)

    doc.add_heading("5.1. Обзорная схема", level=2)
    add_code_block(doc, (
        "persons (1) ──┬── (N) person_identification_keys\n"
        "              ├── (N) person_external_ids\n"
        "              ├── (N) person_defects\n"
        "              ├── (N) person_documents\n"
        "              ├── (N) person_deferred_cessations\n"
        "              ├── (N) person_review_queue\n"
        "              └── (N) person_review_history\n"
        "\n"
        "ext_persons (1) ──┬── (N) ext_person_defects\n"
        "                  ├── (N) ext_person_cessations\n"
        "                  └── (N) ext_person_deferred_cessations"
    ))

    doc.add_heading("5.2. Золотые таблицы", level=2)
    add_table(doc,
        ["Таблица", "Назначение"],
        [
            ["persons", "Единая запись лица (только PersonID + таймстемпы, без ПДн)"],
            ["person_identification_keys", "HMAC-ключи для сопоставления лиц"],
            ["person_external_ids", "Связи с внешними информационными системами"],
            ["person_defects", "Дефекты данных (invalid_inn, invalid_snils, dul_incomplete)"],
            ["person_documents", "Документы (ДУЛ)"],
            ["person_deferred_cessations", "Записи отложенной прекращения обработки ПДн"],
            ["person_review_queue", "Очередь на ручную обработку (Ambiguous)"],
            ["person_review_history", "История разрешённых конфликтов"],
        ],
    )

    doc.add_heading("5.3. Staging-таблицы (ext_*)", level=2)
    doc.add_paragraph(
        "Staging-таблицы хранят сырые данные всех запросов resolve для аудита. "
        "Данные ссылаются на ext_persons, а не на persons. "
        "HMAC-ключи stored в staging-таблицах для отслеживания истории."
    )

    # 6. Алгоритм идентификации
    doc.add_heading("6. Алгоритм идентификации", level=1)

    doc.add_heading("6.1. Нормализация данных", level=2)
    add_table(doc,
        ["Поле", "Правила нормализации"],
        [
            ["ФИО", "Trim → схлопывание пробелов → Unicode NFC → ToUpperInvariant"],
            ["ИНН", "Trim → схлопывание → NFC → ToUpperInvariant + валидация (10/12 цифр, контрольная сумма)"],
            ["СНИЛС", "Те же правила + валидация (11 цифр, контрольная сумма)"],
            ["ДУЛ", "Тип: нормализация как ФИО. Серия/номер: trim + collapse. Объединение: type|series|number"],
        ],
    )

    doc.add_heading("6.2. Вычисление HMAC-SHA256 ключей", level=2)
    doc.add_paragraph("Формируются шесть типов ключей:")
    add_table(doc,
        ["Тип ключа", "Формула"],
        [
            ["inn", "HMAC(normalized_inn)"],
            ["snils", "HMAC(normalized_snils)"],
            ["dul", "HMAC(type|series|number)"],
            ["inn_fio", "HMAC(normalized_inn|fio)"],
            ["snils_fio", "HMAC(normalized_snils|fio)"],
            ["dul_fio", "HMAC(normalized_dul|fio)"],
        ],
    )
    doc.add_paragraph(
        "Стandalone-ключи fio и fio_full не создаются — ФИО не являются "
        "сильным доказательством идентичности."
    )

    doc.add_heading("6.3. Поиск совпадений и scoring", level=2)
    doc.add_paragraph("Алгоритм определения статуса:")
    add_numbered(doc, "Вычислить все доступные ключи из запроса")
    add_numbered(doc, "Найти все person_identification_keys с совпадающими key_value")
    add_numbered(doc, "Если найдено ровно 1 лицо → Matched")
    add_numbered(doc, "Если 0 лиц → Unmatched (создание нового)")
    add_numbered(doc, "Если >1 лица: есть ИНН, совпадающий ровно с 1 лицом → Auto-merge")
    add_numbered(doc, "Иначе → Ambiguous (конфликт, ручной разбор)")

    doc.add_heading("6.4. Автоматическое слияние (auto-merge)", level=2)
    doc.add_paragraph(
        "Когда при resolve с ИНН совпадение найдено ровно у одного лица — "
        "все ключи, внешние ID и документы сливаемого лица переносятся "
        "в выживающее через PersonMergeService."
    )

    doc.add_heading("6.5. Автозакрытие конфликтов", level=2)
    doc.add_paragraph(
        "При последующем resolve для того же externalPersonId с исправленными "
        "данными → Ambiguous закрывается автоматически. "
        "История сохраняется в person_review_history."
    )

    # 7. Прекращение обработки ПДн
    doc.add_heading("7. Прекращение обработки ПДн", level=1)

    doc.add_heading("7.1. Двухфазный подход", level=2)
    add_table(doc,
        ["Фаза", "Метод", "Действие"],
        [
            ["1", "CeaseProcessingAsync", "Пометить внешние ключи (processing_status = 'cessation')"],
            ["2", "ReconcileAsync", "Удалить помеченные staging-записи и золотые записи (если нет ссылок)"],
        ],
    )

    doc.add_heading("7.2. Мгновенный отзыв", level=2)
    doc.add_paragraph("POST /persons/cessation — немедленное удаление данных по указанным связям.")

    doc.add_heading("7.3. Отложенный отзыв", level=2)
    doc.add_paragraph(
        "POST /persons/cessation/deferred — запланированное удаление на указанную дату. "
        "Если до наступления даты данные будут добавлены снова — отзыв автоматически отменяется."
    )

    doc.add_heading("7.4. Реконсилизация (Worker)", level=2)
    doc.add_paragraph(
        "Worker выполняет реконсилизацию по cron (по умолчанию каждый час):"
    )
    add_numbered(doc, "ProcessDeferredCessationsAsync — преобразует отложенные отзывы в немедленные")
    add_numbered(doc, "ReconcileAsync — удаляет помеченные записи и orphaned золотые записи")

    doc.add_heading("7.5. Порядок удаления (FK-безопасность)", level=2)
    add_table(doc,
        ["#", "Таблица", "FK"],
        [
            ["1", "person_identification_keys", "→ persons(id) RESTRICT"],
            ["2", "person_defects", "→ persons(id) RESTRICT"],
            ["3", "person_documents", "→ persons(id) RESTRICT"],
            ["4", "person_deferred_cessations", "→ persons(id) RESTRICT"],
            ["5", "ext_person_defects", "→ ext_persons(id) RESTRICT"],
            ["6", "ext_persons", "→ persons(id) SET NULL"],
            ["7", "persons", "(корневая таблица)"],
        ],
    )

    # 8. API сервиса
    doc.add_heading("8. API сервиса", level=1)
    doc.add_paragraph("Base URL: http://<host>:5000")
    doc.add_paragraph("Swagger UI: http://<host>:5000/swagger")

    add_table(doc,
        ["Метод", "Путь", "Назначение"],
        [
            ["POST", "/persons/resolve", "Идентификация лица (поиск или создание)"],
            ["GET", "/persons/{masterId}", "Получение данных лица (без ПДн)"],
            ["POST", "/persons/{masterId}/identifiers", "Добавление связи с внешней системой"],
            ["POST", "/persons/validate/inn", "Валидация ИНН"],
            ["POST", "/persons/validate/snils", "Валидация СНИЛС"],
            ["POST", "/persons/cessation", "Прекращение обработки ПДн (немедленное)"],
            ["POST", "/persons/cessation/deferred", "Отложенное прекращение обработки ПДн"],
            ["POST", "/persons/cessation/reconcile", "Реконсилизация (удаление помеченных записей)"],
            ["GET", "/persons/dul-classifier", "Справочник видов ДУЛ"],
            ["GET", "/persons/review", "Очередь на ручную обработку"],
        ],
    )

    doc.add_heading("8.1. POST /persons/resolve — Пример запроса", level=2)
    add_code_block(doc, (
        '{\n'
        '  "lastName": "Иванов",\n'
        '  "firstName": "Иван",\n'
        '  "middleName": "Иванович",\n'
        '  "evidence": {\n'
        '    "inn": "7707083893",\n'
        '    "snils": "12345678964",\n'
        '    "dulType": "21",\n'
        '    "dulSeries": "4510",\n'
        '    "dulNumber": "123456"\n'
        '  },\n'
        '  "identifiers": [\n'
        '    {\n'
        '      "sourceSystemId": "CRM",\n'
        '      "externalPersonId": "ext-12345",\n'
        '      "externalPersonType": "client"\n'
        '    }\n'
        '  ]\n'
        '}'
    ))

    doc.add_heading("8.2. POST /persons/resolve — Пример ответа", level=2)
    add_code_block(doc, (
        '{\n'
        '  "status": "Matched",\n'
        '  "personId": "a1b2c3d4-e5f6-7890-abcd-ef1234567890"\n'
        '}'
    ))
    doc.add_paragraph("Возможные статусы: Matched, Unmatched, Ambiguous.")

    # 9. Требования к производительности
    doc.add_heading("9. Требования к производительности", level=1)
    add_table(doc,
        ["Метрика", "Целевое значение"],
        [
            ["Latency API (p95)", "< 200 ms"],
            ["Latency API (p99)", "< 500 ms"],
            ["Concurrent Users", "1000"],
            ["Data Availability", "99.9%"],
        ],
    )

    # 10. Безопасность
    doc.add_heading("10. Безопасность", level=1)
    add_bullet(doc, "HMAC-SHA256 хеширование: необратимое преобразование с секретным ключом")
    add_bullet(doc, "ПДн не хранятся в золотых записях (persons) — только хеши и таймстемпы")
    add_bullet(doc, "Секретный ключ HMAC хранится в .env (не в git)")
    add_bullet(doc, "Аудит IP-адресов: каждый запрос фиксирует source_ip в staging-таблицах")
    add_bullet(doc, "Логирование без ПДн: только PersonId, HMAC-хеши, внешние ID")
    add_bullet(doc, "Регулярная проверка CVE: БДУ ФСТЭК, GitHub Security Advisories")
    add_bullet(doc, "Сетевая изоляция: bridge-сеть mnemonios, доступ из корпоративного контура")
    add_bullet(doc, "Транзакционность: все операции записи выполняются в явных транзакциях")

    doc.save("docs/guides/explanatory_note.docx")
    print("  [OK] explanatory_note.docx")


# ─────────────────────── Точка входа ───────────────────────


if __name__ == "__main__":
    print("Генерация документации ЕДИН...")
    create_admin_guide()
    create_user_guide()
    create_explanatory_note()
    print("Готово! Файлы в docs/guides/")
